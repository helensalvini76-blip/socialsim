#!/usr/bin/env python
"""Build the two on-the-day documents.

    python tools/build_day_pack.py "<output folder>"

Produces:
    Exercise Jupiter - MediaSim facilitator run sheet.docx
    Exercise Jupiter - MediaSim briefing for the communications team.docx
"""

import os, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = 'Aptos'
HEAD_FONT = 'Aptos Display'
NAVY = '1F4E79'
AMBER = 'B45309'
GREY = '6B7280'
RED = '9B1C1C'
GREEN = '166534'

BASE = 'https://mediasim.bcrconsultants.co.uk'


# ── helpers ──────────────────────────────────────────────────────────
def shade(cell, fill):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:color'), 'auto')
    el.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(el)


def borders(table, inside='E5E7EB', outside='D1D5DB'):
    b = OxmlElement('w:tblBorders')
    for edge, colour in [('top', outside), ('bottom', outside), ('left', 'auto'),
                         ('right', 'auto'), ('insideH', inside), ('insideV', 'auto')]:
        e = OxmlElement('w:%s' % edge)
        if colour == 'auto':
            e.set(qn('w:val'), 'none')
            e.set(qn('w:sz'), '0')
        else:
            e.set(qn('w:val'), 'single')
            e.set(qn('w:sz'), '4')
            e.set(qn('w:color'), colour)
        e.set(qn('w:space'), '0')
        b.append(e)
    table._tbl.tblPr.append(b)


def run(par, text, size=10.5, bold=False, italic=False, colour=None, font=None):
    r = par.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if colour:
        r.font.color.rgb = RGBColor.from_string(colour)
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.insert(0, rf)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rf.set(qn(attr), font or BODY_FONT)
    return r


def tight(par, before=0, after=3):
    par.paragraph_format.space_before = Pt(before)
    par.paragraph_format.space_after = Pt(after)
    return par


def new_doc(margin=1.9):
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = BODY_FONT
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
    sec = doc.sections[0]
    for a in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(sec, a, Cm(margin))
    return doc, sec.page_width - sec.left_margin - sec.right_margin


def title(doc, main, sub):
    run(tight(doc.add_paragraph(), after=0), main, size=25, bold=True, colour=NAVY, font=HEAD_FONT)
    run(tight(doc.add_paragraph(), after=11), sub, size=12.5, colour='4B5563')


def heading(doc, text, size=14):
    run(tight(doc.add_paragraph(), before=11, after=4), text, size=size, bold=True, colour=NAVY, font=HEAD_FONT)


def para(doc, *runs, **kw):
    p = tight(doc.add_paragraph(), after=kw.get('after', 5))
    for r in runs:
        if isinstance(r, str):
            run(p, r, size=kw.get('size', 10.5))
        else:
            run(p, r[0], size=kw.get('size', 10.5), **r[1])
    return p


def bullets(doc, items, size=10.5):
    for it in items:
        p = tight(doc.add_paragraph(), after=3)
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        run(p, '•   ', size=size, colour=GREY)
        if isinstance(it, tuple):
            run(p, it[0], size=size, bold=True)
            run(p, it[1], size=size)
        else:
            run(p, it, size=size)


def table(doc, width, rows, widths, header=None, sizes=(9.5, 9.5)):
    t = doc.add_table(rows=0, cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    borders(t)
    if header:
        r = t.add_row()
        for i, h in enumerate(header):
            r.cells[i].width = widths[i]
            shade(r.cells[i], 'F1F5F9')
            run(tight(r.cells[i].paragraphs[0]), h, size=9, bold=True, colour=NAVY)
    for row in rows:
        r = t.add_row()
        for i, cell in enumerate(row):
            r.cells[i].width = widths[i]
            bold = (i == 0 and len(row) > 1)
            run(tight(r.cells[i].paragraphs[0]), cell, size=sizes[0] if i == 0 else sizes[1], bold=bold)
    return t


def callout(doc, width, label, text, fill='FFFBEB', colour=AMBER):
    t = doc.add_table(rows=1, cols=1)
    borders(t, inside=fill, outside='FDE68A')
    c = t.rows[0].cells[0]
    c.width = width
    shade(c, fill)
    p = tight(c.paragraphs[0], after=0)
    run(p, label + '   ', size=9.5, bold=True, colour=colour)
    run(p, text, size=9.5, colour=colour)
    tight(doc.add_paragraph(), after=4)


# ── run sheet ────────────────────────────────────────────────────────
def run_sheet(folder):
    doc, W = new_doc()
    title(doc, 'Exercise Jupiter', 'MediaSim — facilitator run sheet')

    table(doc, W, [
        ['Exercise', 'Jupiter — fire and evacuation, The Kirkwood, Dalton site'],
        ['Date', '10 September 2026'],
        ['Exercise play', '13:00 – 15:30  (T+0 to T+150)'],
        ['Exercise Director', 'Helen Salvini'],
        ['Your screen', BASE + '/dashboard.html'],
        ['Comms team', BASE + '/jupiter.html'],
        ['ICC projector', BASE + '/wall.html'],
        ['Dashboard passphrase', 'jupiter2026'],
    ], [Cm(4.4), W - Cm(4.4)])

    callout(doc, W, 'BEFORE ANYTHING ELSE',
            'Run a full dry run first using ?session=dryrun on the end of all three links. '
            'That keeps the real session clean for the 10th.')

    heading(doc, 'The week before')
    bullets(doc, [
        ('Dry run. ', 'Two phones and the laptop, at 30x speed so the whole exercise takes five minutes.'),
        ('Confirm the fire location ', 'and send it over — it is a one-line change.'),
        ('Brief Lynne ', 'that the ex-staff fire door allegation is in the script at T+118, and that a version can also be fired earlier by hand.'),
        ('Check the signal ', 'in the car park and the ICC. The tool survives a dropout, but you want to know in advance.'),
        ('Send the comms team ', 'their link and the one-page briefing the day before, so nobody is learning the tool during the exercise.'),
    ])

    heading(doc, '12:00 – 12:30  Setting up')
    table(doc, W, [
        ['1', 'Open the dashboard on the laptop. Enter the passphrase, or use the link with ?key=jupiter2026 built in.'],
        ['2', 'Check the build number at the top right matches what you were told. A lower number means a stale cache — close the tab completely and reopen.'],
        ['3', 'Press RESET once. This clears anything left from the dry run on every device. It will ask you to confirm.'],
        ['4', 'Confirm the clock reads T+000 13:00 and the button says START EXERCISE. Nothing fires until you press it.'],
        ['5', 'Open the wall display on the projector. It should show the holding screen.'],
        ['6', 'Open the participant link on one phone and check a post appears. Then close it.'],
        ['7', 'Confirm the dashboard says "live - facilitator" with a green dot.'],
    ], [Cm(1.0), W - Cm(1.0)], sizes=(9.5, 9.5))

    callout(doc, W, 'DO NOT',
            'open the dashboard on two machines. Both will try to drive the clock. A red banner appears if it happens.')

    heading(doc, '12:45  Comms team')
    bullets(doc, [
        'Hand out the participant link. Anyone can join at any point and will land at the right moment.',
        'Two minutes of orientation is enough: five tabs along the bottom, tap to post, tap Reply on anything.',
        'Remind them nothing is real and nothing leaves the room.',
    ])

    heading(doc, '13:00  Start')
    para(doc, ('Press the green ', {}), ('START EXERCISE', {'bold': True}),
         (' button. Confirm when prompted. The clock begins at T+0 and every inject then fires by itself.', {}))

    heading(doc, 'During the exercise')
    para(doc, 'It runs itself. You do not need to trigger anything. Intervene only when you want to:', size=10.5)
    table(doc, W, [
        ['Quick-fire', 'One click sends an extra inject. Eight groups: Media, Families, Rumour, Staff, Partners, Positive, Internal, Enquiries. Hover to read the full text first.'],
        ['now', 'On any timeline row, pulls that scripted inject forward to this moment.'],
        ['Pause', 'Stops the clock for everyone. Use for NO DUFF, EXERCISE STOP, or a break.'],
        ['Speed', 'Only if you are running behind. Leave at 1x normally.'],
        ['Watermark', 'Puts EXERCISE across every post on every device, live. Use if anyone looks uneasy about screenshots.'],
    ], [Cm(2.6), W - Cm(2.6)])

    para(doc, ('Watch three things: ', {'bold': True}),
         ('the middle column shows everything the comms team posts and how long they took to reply; the indicators show '
          'time to first holding message against your 30-minute target; and the unanswered enquiries panel shows what '
          'they have let slide.', {}), after=6)

    heading(doc, 'Moments worth watching')
    table(doc, W, [
        ['T+15', 'First media approach. The 30-minute clock is running.'],
        ['T+17', 'First family asking publicly. Karen Whitfield, father on the IPU.'],
        ['T+20', 'First formal enquiry lands in the inbox with a 30-minute deadline.'],
        ['T+25', 'Staff speculation in the group — and the same rumour on Facebook.'],
        ['T+35', 'The false death rumour enters. Watch whether they correct it.'],
        ['T+55', 'Broadcast interview request with a hard 17:30 deadline.'],
        ['T+58', 'Reception asks for a line she is allowed to use. The sharpest internal test.'],
        ['T+70', 'Late shift asking whether to come in.'],
        ['T+95', 'Two relatives arrive at the gate.'],
        ['T+118', 'Ex-staff fire door allegation.'],
        ['T+130', 'ICB asks for a SitRep.'],
    ], [Cm(1.7), W - Cm(1.7)])

    heading(doc, '15:30  End')
    bullets(doc, [
        ('Pause the clock ', 'when the Exercise Director announces the end.'),
        ('Press Export debrief. ', 'A text file downloads with the indicators, everything the comms team said with timings, which enquiries were answered and how long each took, which were never answered, and the full inject log.'),
        ('Do not press Reset ', 'until you have exported. Reset clears everything for everyone.'),
    ])

    heading(doc, 'Hot debrief — what the export gives you')
    bullets(doc, [
        'Time to first holding message, against the 30-minute indicator in your pack.',
        'Every statement they issued, with the exact minute.',
        'How long each enquiry waited before it was answered.',
        'Which enquiries were never answered at all — usually the most uncomfortable and most useful finding.',
    ])

    heading(doc, 'If something goes wrong')
    table(doc, W, [
        ['Blank screen', 'The page now shows an error rather than nothing. Close the tab completely and reopen the link.'],
        ['Old content', 'Check the build number. A lower number is a stale cache — close the tab fully, do not just refresh.'],
        ['Someone lost signal', 'Their screen shows a "No signal" pill. Anything they wrote is queued and sends when they reconnect.'],
        ['Clock jumped back', 'Nothing is lost. Posts reappear as the clock passes their time again.'],
        ['Laptop restarted', 'Reopen the dashboard. It rejoins the exercise at the right time rather than restarting it.'],
        ['Red banner', 'Another dashboard is driving. Close one of them.'],
    ], [Cm(3.4), W - Cm(3.4)])

    out = os.path.join(folder, 'Exercise Jupiter - MediaSim facilitator run sheet.docx')
    doc.save(out)
    return out


# ── comms team briefing ──────────────────────────────────────────────
def briefing(folder):
    doc, W = new_doc(margin=1.5)
    run(tight(doc.add_paragraph(), after=0), 'Exercise Jupiter', size=21, bold=True, colour=NAVY, font=HEAD_FONT)
    run(tight(doc.add_paragraph(), after=8), 'MediaSim — briefing for the communications team', size=11.5, colour='4B5563')

    para(doc, ('You will be using a simulated social media environment during this exercise. Everything in it is '
               'invented, and none of it is visible outside the exercise.', {}), after=6, size=10)

    heading(doc, 'Your link', size=12.5)
    p = tight(doc.add_paragraph(), after=5)
    run(p, BASE + '/jupiter.html', size=12.5, bold=True, colour=NAVY)
    para(doc, 'Phone, tablet or laptop. You can join at any point and will arrive at the right moment. '
              'If you close it by accident, just open it again.', after=7, size=10)

    heading(doc, 'What you will see', size=12.5)
    table(doc, W, [
        ['Facebook', 'Where most of the local reaction lives. Families and the public.'],
        ['X', 'Media, partner agencies and faster-moving speculation.'],
        ['Instagram', 'Photos from the scene.'],
        ['Staff', 'A group chat with colleagues across the hospice.'],
        ['Enquiries', 'Emails, phone messages and reception notes that need a response.'],
    ], [Cm(2.6), W - Cm(2.6)])

    heading(doc, 'What you can do', size=12.5)
    bullets(doc, [
        ('Post as The Kirkwood. ', 'The pencil button, or Post on a laptop. Pick the platform first by tapping its tab.'),
        ('Reply to anything. ', 'Every post has a reply control. Your reply appears as The Kirkwood, badged OFFICIAL.'),
        ('Reply to individual comments. ', 'Open a thread with "View all comments" and answer one specific person.'),
        ('Respond to enquiries. ', 'Each shows "Awaiting response" until answered. Your response appears underneath.'),
    ], size=10)

    callout(doc, W, 'PEOPLE ANSWER BACK.',
            'This is not a static screen. If you issue a statement, people respond to it — some warmly, some not. '
            'Journalists will push for numbers. Relatives will ask what happens next. Treat it as you would the real thing.',
            fill='EFF6FF', colour='1E40AF')

    heading(doc, 'What this is not')
    bullets(doc, [
        'It is not a test of you as an individual. It is a test of the arrangements.',
        'There are no trick questions and no scoring of your writing.',
        'Nothing you post can reach the outside world. It exists only inside the exercise.',
    ])

    heading(doc, 'Ground rules', size=12.5)
    bullets(doc, [
        ('Do not screenshot or share anything. ', 'The content is deliberately realistic and would alarm people out of context.'),
        ('Treat the information as real ', 'within the scenario, but take no action outside it — no real calls, emails or posts.'),
        ('Say NO DUFF ', 'if anything real happens. Say EXERCISE STOP for a safety concern.'),
        ('If your screen says "No signal", ', 'keep going. Anything you write is saved and sends when you reconnect.'),
        ('This is a test of the arrangements, ', 'not of you. Nothing you write is scored.'),
    ], size=10)

    heading(doc, 'A note on the content', size=12.5)
    para(doc, 'Some of what appears will be uncomfortable — inaccurate claims, criticism of the hospice, distressed '
              'relatives. That is deliberate, because it is what happens. Every person shown is invented; no real '
              'journalist, resident, relative or colleague is named or depicted. If anything troubles you during or '
              'after the exercise, say so to the Exercise Director.', after=4, size=10)

    out = os.path.join(folder, 'Exercise Jupiter - MediaSim briefing for the communications team.docx')
    doc.save(out)
    return out


if __name__ == '__main__':
    folder = sys.argv[1] if len(sys.argv) > 1 else '.'
    for f in (run_sheet(folder), briefing(folder)):
        print('wrote %s  (%.0fKB)' % (os.path.basename(f), os.path.getsize(f) / 1024))
