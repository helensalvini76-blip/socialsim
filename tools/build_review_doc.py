#!/usr/bin/env python
"""Build the Exercise Jupiter content review document.

Regenerate after any content change:

    python tools/extract_content.py tools/jupiter-content.json
    python tools/build_review_doc.py tools/jupiter-content.json "out.docx"
"""

import io, json, os, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
BODY_FONT = 'Aptos'
HEAD_FONT = 'Aptos Display'

PLATFORM = {'x': 'X', 'fb': 'Facebook', 'ig': 'Instagram',
            'staff': 'STAFF GROUP', 'inbox': 'ENQUIRY'}

PHASES = [
    (-999, -1,  'Before the alarm',
     'Ordinary hospice life. Seeds the feed so it does not start empty - participants can scroll back into a normal day.'),
    (0, 24,     'Phase 1 - Alarm and escalation',
     'T+0 to T+25. First media approach at T+15; first family enquiry at T+17.'),
    (25, 46,    'Phase 2 - Live movement',
     'T+20 to T+65. The false death rumour enters at T+35.'),
    (47, 67,    'Phase 3 - Clearing station and families',
     'T+45 to T+85. Rumour peaks, the reporter goes direct to families and staff, and the rain arrives.'),
    (68, 100,   'Phase 4 - Dispersal and media peak',
     'T+65 to T+115. Pressure on the comms team is highest here.'),
    (101, 999,  'Phase 5 - Recovery',
     'T+100 to T+150. The story turns from incident to scrutiny. The ex-staff fire door allegation lands at T+118.'),
]


def clock(minute):
    total = 13 * 60 + minute
    h = (total % 1440) // 60
    m = total % 60
    return '%02d:%02d' % (h, m)


def stamp(minute):
    label = ('T%d' % minute) if minute < 0 else ('T+%d' % minute)
    return '%-7s %s' % (label, clock(minute))


def shade(cell, hex_fill):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:color'), 'auto')
    el.set(qn('w:fill'), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def no_borders(table, inside='F1F2F4', outside='E3E5E8'):
    tbl = table._tbl
    pr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    spec = [('top', outside), ('bottom', outside), ('left', 'auto'),
            ('right', 'auto'), ('insideH', inside), ('insideV', 'auto')]
    for edge, colour in spec:
        e = OxmlElement('w:%s' % edge)
        if colour == 'auto':
            e.set(qn('w:val'), 'none')
            e.set(qn('w:sz'), '0')
        else:
            e.set(qn('w:val'), 'single')
            e.set(qn('w:sz'), '4')
            e.set(qn('w:color'), colour)
        e.set(qn('w:space'), '0')
        borders.append(e)
    pr.append(borders)


def run(par, text, size=10.5, bold=False, italic=False, colour=None, font=None, space_after=None):
    r = par.add_run(text)
    r.font.name = font or BODY_FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if colour:
        r.font.color.rgb = RGBColor.from_string(colour)
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.insert(0, rf)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rf.set(qn(attr), font or BODY_FONT)
    return r


def tight(par, before=0, after=2):
    par.paragraph_format.space_before = Pt(before)
    par.paragraph_format.space_after = Pt(after)
    return par


def main():
    src = sys.argv[1] if len(sys.argv) > 1 else os.path.join(ROOT, 'tools', 'jupiter-content.json')
    out = sys.argv[2] if len(sys.argv) > 2 else 'Exercise Jupiter - MediaSim content for review.docx'
    data = json.load(io.open(src, encoding='utf-8'))
    people = data['people']

    def who(key):
        p = people.get(key)
        return ('%s  %s' % (p['name'], p['handle'])) if p else key

    doc = Document()
    st = doc.styles['Normal']
    st.font.name = BODY_FONT
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)

    sec = doc.sections[0]
    for attr in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(sec, attr, Cm(1.9))
    usable = sec.page_width - sec.left_margin - sec.right_margin

    # ── Title ────────────────────────────────────────────────────
    tight(doc.add_paragraph(), after=0)
    p = doc.paragraphs[-1]
    run(p, 'Exercise Jupiter', size=28, bold=True, colour='1F4E79', font=HEAD_FONT)
    p2 = tight(doc.add_paragraph(), after=10)
    run(p2, 'MediaSim content script — for review', size=13, colour='4B5563')

    allposts = sorted(data['baseline'] + data['script'], key=lambda x: x['min'])
    nthreads = len(data['threads'])
    ncomments = sum(len(v) for v in data['threads'].values())

    facts = [
        ('Exercise', 'Jupiter — fire and evacuation, The Kirkwood, Dalton site'),
        ('Date', '10 September 2026, 13:00 – 15:30  (T+0 to T+150)'),
        ('Posts', '%d  (%d before the alarm, %d during)' % (len(allposts), len(data['baseline']), len(data['script']))),
        ('Comments', '%d threads, %d comments' % (nthreads, ncomments)),
        ('Direct enquiries', '%d — tracked for answered / unanswered reporting' % sum(1 for x in allposts if x['enquiry'])),
        ('Accounts', '%d — every individual invented' % len(people)),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_borders(t)
    for k, v in facts:
        row = t.add_row()
        row.cells[0].width = Cm(4.2)
        row.cells[1].width = usable - Cm(4.2)
        shade(row.cells[0], 'F7F8FA')
        run(tight(row.cells[0].paragraphs[0]), k, size=9.5, bold=True)
        run(tight(row.cells[1].paragraphs[0]), v, size=9.5)

    doc.add_paragraph()
    p = tight(doc.add_paragraph(), after=6)
    run(p, 'How to use this document.  ', size=9.5, bold=True)
    run(p, 'Everything below is the exact content the comms team will see, in the order they will see it. '
           'Comment on anything that should change — wording, tone, timing, or anything that would not ring true '
           'to a Kirkwood audience — and it will be edited in the tool. Amber notes are facilitator-only and are '
           'never shown to participants.', size=9.5)
    p = tight(doc.add_paragraph(), after=6)
    run(p, 'Identities.  ', size=9.5, bold=True)
    run(p, 'The Kirkwood, the Dalton site and the local geography are real. Every individual, journalist, resident, '
           'relative and news outlet named here is invented.', size=9.5)

    # ── Phases ───────────────────────────────────────────────────
    for lo, hi, name, note in PHASES:
        posts = [x for x in allposts if lo <= x['min'] <= hi]
        if not posts:
            continue
        doc.add_page_break()
        h = tight(doc.add_paragraph(), after=2)
        run(h, name, size=16, bold=True, colour='1F4E79', font=HEAD_FONT)
        hp = tight(doc.add_paragraph(), after=8)
        run(hp, note, size=9.5, italic=True, colour='6B7280')

        for post in posts:
            thread = data['threads'].get(str(post['min']), [])
            tbl = doc.add_table(rows=0, cols=2)
            no_borders(tbl)
            wl, wr = Cm(2.9), usable - Cm(2.9)

            r = tbl.add_row()
            r.cells[0].width, r.cells[1].width = wl, wr
            shade(r.cells[0], 'F7F8FA')
            run(tight(r.cells[0].paragraphs[0]), stamp(post['min']), size=9, bold=True, font='Consolas')

            c = r.cells[1]
            meta = tight(c.paragraphs[0], after=1)
            plat = post['plat']
            colour = '0F7B3F' if plat == 'staff' else ('6B21A8' if plat == 'inbox' else '1F4E79')
            run(meta, PLATFORM.get(plat, plat), size=8.5, bold=True, colour=colour)
            if post.get('via'):
                run(meta, '   •   ' + post['via'], size=8.5, colour='6B7280')
            if post.get('packRef'):
                run(meta, '   •   PACK ' + post['packRef'], size=8.5, bold=True, colour='B45309')
            if post.get('enquiry'):
                run(meta, '   •   DIRECT ENQUIRY', size=8.5, bold=True, colour='7C2D12')
            if post.get('img'):
                run(meta, '   •   photo', size=8.5, colour='6B7280')
            run(tight(c.add_paragraph(), after=2), who(post['who']), size=10, bold=True)
            if post.get('subject'):
                run(tight(c.add_paragraph(), after=2), post['subject'], size=10.5, bold=True)
            run(tight(c.add_paragraph(), after=3), post['text'], size=10.5)

            for cm in thread:
                r = tbl.add_row()
                r.cells[0].width, r.cells[1].width = wl, wr
                shade(r.cells[0], 'F7F8FA')
                cc = r.cells[1]
                hp2 = tight(cc.paragraphs[0], after=1)
                run(hp2, '↳  ', size=9.5, colour='9CA3AF')
                run(hp2, who(cm['who']), size=9, bold=True, colour='374151')
                body = tight(cc.add_paragraph(), after=2)
                body.paragraph_format.left_indent = Cm(0.5)
                run(body, cm['text'], size=9.5, colour='374151')

            if post.get('note'):
                r = tbl.add_row()
                r.cells[0].width, r.cells[1].width = wl, wr
                shade(r.cells[0], 'F7F8FA')
                shade(r.cells[1], 'FFFBEB')
                np_ = tight(r.cells[1].paragraphs[0])
                run(np_, 'Facilitator note   ', size=8.5, bold=True, colour='7C2D12')
                run(np_, post['note'], size=8.5, italic=True, colour='7C2D12')

            tight(doc.add_paragraph(), after=4)

    doc.save(out)
    print('wrote %s  (%.0fKB)' % (out, os.path.getsize(out) / 1024))
    print('posts %d  threads %d  comments %d' % (len(allposts), nthreads, ncomments))


if __name__ == '__main__':
    main()
